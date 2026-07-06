import sys, re, glob, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_runs(p, text):
    # split on **bold** and `code`
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.color.rgb = RGBColor(0xC0,0x30,0x30)
        else:
            p.add_run(part)

def is_table_sep(line):
    s = line.strip().strip('|')
    return bool(s) and set(s.replace('|','').replace(' ','')) <= set('-:')

def convert(md_path, out_path):
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
    lines = open(md_path, encoding='utf-8').read().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # table block
        if line.strip().startswith('|') and i+1 < len(lines) and is_table_sep(lines[i+1]):
            rows = []
            header = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
            for j,h in enumerate(header):
                cell = t.rows[0].cells[j]; cell.paragraphs[0].clear()
                add_runs(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs: r.bold = True
            for row in rows:
                cells = t.add_row().cells
                for j,val in enumerate(row):
                    if j < len(cells):
                        cells[j].paragraphs[0].clear(); add_runs(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue
        if not line.strip():
            i += 1; continue
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            lvl = min(len(m.group(1)), 4)
            h = doc.add_heading('', level=lvl); add_runs(h, re.sub(r'[*`]','',m.group(2)))
        elif line.strip().startswith('>'):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, line.strip()[1:].strip())
            for r in p.runs: r.italic = True
        elif re.match(r'^\s*[-*]\s+', line):
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, re.sub(r'^\s*[-*]\s+','',line))
        elif re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\s*\d+\.\s+','',line))
        elif line.strip() in ('---','***','___'):
            pass
        else:
            p = doc.add_paragraph(); add_runs(p, line)
        i += 1
    doc.save(out_path)
    return out_path

folder = sys.argv[1]
done = []
for md in sorted(glob.glob(folder + '/*.md')):
    out = md[:-3] + '.docx'
    try:
        convert(md, out); done.append(os.path.basename(out))
    except Exception as e:
        print('ERR', os.path.basename(md), e)
for d in done: print('  ->', d)
print(f'{len(done)} archivos Word creados')
